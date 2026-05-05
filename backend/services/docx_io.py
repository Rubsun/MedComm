"""Экспорт/импорт уроков в .docx.

На текущем этапе поддерживается только text/image/video.
Блоки practice/quiz экспортируются как плейсхолдер; импорт их пропускает,
сохраняя существующие данные в БД.

Структура файла:

    Заголовок урока (Heading 1)
    [описание урока]
    === BLOCK 0 | text ===          (Heading 2 — маркер блока)
    <параграфы текста с форматированием>
    === BLOCK 1 | image ===
    URL: https://...
    ALT: ...
    === BLOCK 2 | practice ===
    [не редактируется через .docx]
"""

from __future__ import annotations

import io
import re
from typing import Iterable

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.document import Document as _Document
from docx.text.paragraph import Paragraph

BLOCK_HEADING_RE = re.compile(r"^===\s*BLOCK\s+(\d+)\s*\|\s*(\w+)\s*===\s*$")
EDITABLE_TYPES = {"text", "image", "video"}
PLACEHOLDER_TYPES = {"practice", "quiz"}
SUPPORTED_TYPES = EDITABLE_TYPES | PLACEHOLDER_TYPES

INLINE_FORMATS = {
    "strong": "bold",
    "b": "bold",
    "em": "italic",
    "i": "italic",
    "u": "underline",
}


# ---------- HTML -> docx ----------

def _emit_runs(paragraph, node, fmt: dict[str, bool]) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        if fmt.get("bold"):
            run.bold = True
        if fmt.get("italic"):
            run.italic = True
        if fmt.get("underline"):
            run.underline = True
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name == "br":
        paragraph.add_run().add_break()
        return

    next_fmt = dict(fmt)
    if name in INLINE_FORMATS:
        next_fmt[INLINE_FORMATS[name]] = True

    for child in node.children:
        _emit_runs(paragraph, child, next_fmt)


def _children(root) -> Iterable:
    return list(root.children) if hasattr(root, "children") else []


def html_to_paragraphs(doc: _Document, html: str) -> None:
    if not html or not html.strip():
        return
    soup = BeautifulSoup(html, "html.parser")
    body = soup.body or soup

    block_tags = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "blockquote", "pre"}
    has_block = any(
        isinstance(c, Tag) and c.name.lower() in block_tags
        for c in _children(body)
    )

    if not has_block:
        para = doc.add_paragraph()
        for child in _children(body):
            _emit_runs(para, child, {})
        return

    for child in _children(body):
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                doc.add_paragraph().add_run(text)
            continue
        if not isinstance(child, Tag):
            continue
        name = child.name.lower()
        if name == "p":
            para = doc.add_paragraph()
            for sub in child.children:
                _emit_runs(para, sub, {})
        elif name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            # Heading 2 зарезервирован под маркеры блоков — внутренний контент уезжает в 3..5
            level_in = int(name[1])
            level = max(3, min(5, level_in + 1))
            para = doc.add_paragraph(style=f"Heading {level}")
            for sub in child.children:
                _emit_runs(para, sub, {})
        elif name in {"ul", "ol"}:
            list_style = "List Bullet" if name == "ul" else "List Number"
            for li in child.find_all("li", recursive=False):
                para = doc.add_paragraph(style=list_style)
                for sub in li.children:
                    _emit_runs(para, sub, {})
        elif name == "blockquote":
            para = doc.add_paragraph(style="Intense Quote")
            for sub in child.children:
                _emit_runs(para, sub, {})
        elif name == "pre":
            doc.add_paragraph(child.get_text(""))
        else:
            para = doc.add_paragraph()
            for sub in child.children:
                _emit_runs(para, sub, {})


# ---------- docx -> HTML ----------

def _escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _run_to_html(run) -> str:
    text = run.text
    if not text:
        return ""
    text = _escape(text)
    if run.bold:
        text = f"<strong>{text}</strong>"
    if run.italic:
        text = f"<em>{text}</em>"
    if run.underline:
        text = f"<u>{text}</u>"
    return text


def _paragraph_to_inline_html(para: Paragraph) -> str:
    return "".join(_run_to_html(r) for r in para.runs)


def paragraphs_to_html(paragraphs: list[Paragraph]) -> str:
    parts: list[str] = []
    list_buffer: list[str] = []
    list_kind: str | None = None  # 'ul' | 'ol'

    def flush_list() -> None:
        nonlocal list_kind
        if not list_buffer:
            list_kind = None
            return
        tag = list_kind or "ul"
        items = "".join(f"<li>{x}</li>" for x in list_buffer)
        parts.append(f"<{tag}>{items}</{tag}>")
        list_buffer.clear()
        list_kind = None

    for p in paragraphs:
        style = (p.style.name if p.style else "") or ""
        inner = _paragraph_to_inline_html(p)
        is_blank = not inner.strip()

        if style.startswith("List Bullet"):
            if list_kind not in (None, "ul"):
                flush_list()
            list_kind = "ul"
            list_buffer.append(inner)
            continue
        if style.startswith("List Number"):
            if list_kind not in (None, "ol"):
                flush_list()
            list_kind = "ol"
            list_buffer.append(inner)
            continue

        flush_list()
        if is_blank:
            continue

        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 3
            level = max(1, min(6, level))
            parts.append(f"<h{level}>{inner}</h{level}>")
        elif style.startswith("Intense Quote") or style.startswith("Quote"):
            parts.append(f"<blockquote>{inner}</blockquote>")
        else:
            parts.append(f"<p>{inner}</p>")
    flush_list()
    return "".join(parts)


# ---------- export ----------

def lesson_to_docx(lesson_title: str, lesson_description: str, blocks: list) -> bytes:
    doc = Document()
    doc.add_heading(lesson_title or "Без названия", level=1)
    if lesson_description:
        doc.add_paragraph(lesson_description)

    for b in sorted(blocks, key=lambda x: x.sort_order):
        doc.add_heading(f"=== BLOCK {b.sort_order} | {b.type} ===", level=2)
        data = b.data or {}
        if b.type == "text":
            html = data.get("html") or data.get("content") or data.get("markdown") or ""
            html_to_paragraphs(doc, html)
        elif b.type == "image":
            doc.add_paragraph(f"URL: {data.get('url', '')}")
            if data.get("alt"):
                doc.add_paragraph(f"ALT: {data['alt']}")
            if data.get("caption"):
                doc.add_paragraph(f"CAPTION: {data['caption']}")
        elif b.type == "video":
            doc.add_paragraph(f"URL: {data.get('url', '')}")
            if data.get("title"):
                doc.add_paragraph(f"TITLE: {data['title']}")
        elif b.type in PLACEHOLDER_TYPES:
            doc.add_paragraph(
                f"[блок типа '{b.type}' — редактируется только в админке, импорт его не меняет]"
            )
        else:
            doc.add_paragraph(f"[неизвестный тип: {b.type}]")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- import ----------

class DocxImportError(ValueError):
    pass


def _split_into_sections(doc: _Document) -> list[tuple[int, str, list[Paragraph]]]:
    sections: list[tuple[int, str, list[Paragraph]]] = []
    current: tuple[int, str, list[Paragraph]] | None = None
    for p in doc.paragraphs:
        m = BLOCK_HEADING_RE.match((p.text or "").strip())
        if m:
            if current is not None:
                sections.append(current)
            current = (int(m.group(1)), m.group(2).lower(), [])
            continue
        if current is not None:
            current[2].append(p)
    if current is not None:
        sections.append(current)
    return sections


def _kv_lines(paragraphs: list[Paragraph]) -> dict[str, str]:
    out: dict[str, str] = {}
    for p in paragraphs:
        text = (p.text or "").strip()
        if not text or ":" not in text:
            continue
        key, _, value = text.partition(":")
        out[key.strip().upper()] = value.strip()
    return out


def parse_docx(file_bytes: bytes) -> list[dict]:
    """Парсит .docx и возвращает список блоков:
        [{sort_order, type, data, _editable}]

    Блоки типов practice/quiz возвращаются с _editable=False — их данные брать из БД.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise DocxImportError(f"Не удалось открыть .docx: {e}")

    sections = _split_into_sections(doc)
    if not sections:
        raise DocxImportError(
            "В файле не найдено маркеров '=== BLOCK N | type ===' (Heading 2). "
            "Не меняйте структуру заголовков, выгруженных из системы."
        )

    seen: set[int] = set()
    result: list[dict] = []
    for sort_order, btype, paragraphs in sections:
        if btype not in SUPPORTED_TYPES:
            raise DocxImportError(f"Блок sort_order={sort_order}: неизвестный тип '{btype}'.")
        if sort_order in seen:
            raise DocxImportError(f"Дубликат sort_order={sort_order} в файле.")
        seen.add(sort_order)

        if btype in PLACEHOLDER_TYPES:
            result.append({"sort_order": sort_order, "type": btype, "data": None, "_editable": False})
            continue

        if btype == "text":
            html = paragraphs_to_html(paragraphs)
            result.append({
                "sort_order": sort_order, "type": "text",
                "data": {"html": html}, "_editable": True,
            })
        elif btype == "image":
            kv = _kv_lines(paragraphs)
            if not kv.get("URL"):
                raise DocxImportError(f"Блок image sort_order={sort_order}: нет строки 'URL: ...'")
            data = {"url": kv["URL"]}
            if "ALT" in kv:
                data["alt"] = kv["ALT"]
            if "CAPTION" in kv:
                data["caption"] = kv["CAPTION"]
            result.append({"sort_order": sort_order, "type": "image", "data": data, "_editable": True})
        elif btype == "video":
            kv = _kv_lines(paragraphs)
            if not kv.get("URL"):
                raise DocxImportError(f"Блок video sort_order={sort_order}: нет строки 'URL: ...'")
            data = {"url": kv["URL"]}
            if "TITLE" in kv:
                data["title"] = kv["TITLE"]
            result.append({"sort_order": sort_order, "type": "video", "data": data, "_editable": True})

    return result
